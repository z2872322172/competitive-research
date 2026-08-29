from __future__ import annotations

import re
from dataclasses import dataclass
from hashlib import sha256
from html import escape as html_escape
from io import BytesIO
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.schemas import ReportOut, ReportSectionEvidenceOut


@dataclass(frozen=True)
class ReportExportArtifact:
    content: bytes
    content_type: str
    filename: str
    object_key: str | None = None


SUPPORTED_EXPORT_FORMATS = {"markdown", "pdf", "docx"}


def normalize_export_format(value: str | None) -> str:
    return (value or "markdown").strip().lower()


def build_report_markdown(report: ReportOut) -> str:
    parts: list[str] = []
    for section in sorted(report.sections, key=lambda item: item.order_no):
        part = f"## {section.title}\n\n{section.content_markdown}"
        if section.evidence:
            lines = "\n".join(f"- {format_evidence_line(evidence)}" for evidence in section.evidence)
            part += f"\n\n### 引用来源\n\n{lines}"
        parts.append(part)
    return "\n\n".join(parts)


def build_report_export_filename(report: ReportOut, *, format: str, report_title: str | None = None) -> str:
    safe_title = sanitize_filename_component(report_title or f"report-{report.task_id}-v{report.version}")
    if not safe_title.isascii():
        safe_title = f"report-{report.task_id}-v{report.version}"
    extension = {"markdown": "md", "pdf": "pdf", "docx": "docx"}[format]
    return f"{safe_title}.{extension}"


def render_report_export(report: ReportOut, *, format: str, report_title: str | None = None) -> ReportExportArtifact:
    normalized = normalize_export_format(format)
    if normalized == "markdown":
        markdown = build_report_markdown(report)
        return ReportExportArtifact(
            content=markdown.encode("utf-8"),
            content_type="text/markdown; charset=utf-8",
            filename=build_report_export_filename(report, format="markdown", report_title=report_title),
            object_key=build_report_export_object_key(report, format="markdown"),
        )
    if normalized not in SUPPORTED_EXPORT_FORMATS:
        raise ValueError("unsupported_export_format")

    if normalized == "pdf":
        return ReportExportArtifact(
            content=render_report_pdf(report, report_title=report_title),
            content_type="application/pdf",
            filename=build_report_export_filename(report, format=normalized, report_title=report_title),
        )
    if normalized == "docx":
        return ReportExportArtifact(
            content=render_report_docx(report, report_title=report_title),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=build_report_export_filename(report, format=normalized, report_title=report_title),
        )
    raise ValueError("unsupported_export_format")


def build_report_export_object_key(report: ReportOut, *, format: str) -> str:
    extension = {"markdown": "md", "pdf": "pdf", "docx": "docx"}[format]
    return f"reports/{report.task_id}/v{report.version}/report.{extension}"


def build_artifact_record(report: ReportOut, *, format: str, content: bytes, content_type: str) -> dict[str, object]:
    return {
        "artifact_type": format,
        "object_key": build_report_export_object_key(report, format=format),
        "sha256": sha256(content).hexdigest(),
        "content_type": content_type,
        "size_bytes": len(content),
    }


def render_report_pdf(report: ReportOut, *, report_title: str | None = None) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=report_title or f"Report v{report.version}",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportExportTitle",
        parent=styles["Title"],
        alignment=TA_LEFT,
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        spaceAfter=6,
    )
    meta_style = ParagraphStyle(
        "ReportExportMeta",
        parent=styles["BodyText"],
        fontSize=9,
        leading=12,
        textColor="#555555",
        spaceAfter=8,
    )
    section_style = ParagraphStyle(
        "ReportExportSection",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=18,
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "ReportExportBody",
        parent=styles["BodyText"],
        fontSize=10.5,
        leading=14,
        spaceAfter=4,
    )
    bullet_style = ParagraphStyle(
        "ReportExportBullet",
        parent=body_style,
        leftIndent=10,
        firstLineIndent=0,
        bulletIndent=0,
    )

    story: list = [
        Paragraph(html_escape(report_title or f"Report v{report.version}"), title_style),
        Paragraph(
            html_escape(
                f"Task: {report.task_id} | Version: {report.version} | Citation coverage: {round(report.citation_coverage * 100)}%"
            ),
            meta_style,
        ),
    ]

    if report.generated_at is not None:
        story.append(Paragraph(html_escape(f"Generated at: {report.generated_at.isoformat()}"), meta_style))

    for section in sorted(report.sections, key=lambda item: item.order_no):
        story.append(Spacer(1, 4))
        story.append(Paragraph(html_escape(section.title), section_style))
        for paragraph in iter_markdown_paragraphs(section.content_markdown):
            story.append(Paragraph(html_escape(paragraph), body_style))
        if section.evidence:
            story.append(Paragraph("Section evidence", section_style))
            for evidence in section.evidence:
                story.append(Paragraph(html_escape(format_evidence_line(evidence)), bullet_style, bulletText="•"))

    doc.build(story)
    return buffer.getvalue()


def render_report_docx(report: ReportOut, *, report_title: str | None = None) -> bytes:
    document = Document()
    title = report_title or f"Report v{report.version}"
    title_paragraph = document.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta = document.add_paragraph()
    meta.style = document.styles["Normal"]
    meta_run = meta.add_run(
        f"Task: {report.task_id} | Version: {report.version} | Citation coverage: {round(report.citation_coverage * 100)}%"
    )
    meta_run.font.size = Pt(9)
    if report.generated_at is not None:
        generated = document.add_paragraph()
        generated_run = generated.add_run(f"Generated at: {report.generated_at.isoformat()}")
        generated_run.font.size = Pt(9)

    for section in sorted(report.sections, key=lambda item: item.order_no):
        document.add_heading(section.title, level=1)
        for paragraph in iter_markdown_paragraphs(section.content_markdown):
            document.add_paragraph(paragraph)
        if section.evidence:
            document.add_heading("Section evidence", level=2)
            for evidence in section.evidence:
                document.add_paragraph(format_evidence_line(evidence), style="List Bullet")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def iter_markdown_paragraphs(markdown_text: str) -> Iterable[str]:
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("```"):
            continue
        if line.startswith("#"):
            yield line.lstrip("#").strip()
            continue
        if line.startswith(("- ", "* ")):
            yield line[2:].strip()
            continue
        yield line


def format_evidence_line(evidence: ReportSectionEvidenceOut) -> str:
    source_label = evidence.source_title or evidence.publisher or evidence.source_id
    relation = evidence.relation or "supports"
    claim_count = len(evidence.claim_ids)
    claim_label = "Claim" if claim_count == 1 else "Claims"
    reliability = ""
    if evidence.reliability_score is not None:
        level = evidence.reliability_level or "medium"
        reliability = f" | 可靠性 {round(evidence.reliability_score * 100)}%/{level}"
    snapshot = " | 快照可用" if evidence.snapshot_available else ""
    url = f" | {evidence.source_url}" if evidence.source_url else ""
    return (
        f"{evidence.id} | {source_label} | {round(evidence.quality_score * 100)}%"
        f"{reliability} | {relation} | {claim_count} {claim_label}{snapshot}{url}"
    )


def sanitize_filename_component(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\\\|?*]+', "-", value).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned or "report"
